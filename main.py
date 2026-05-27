import asyncio
import base64
import csv
import io
import json
import logging
import mimetypes
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

_bcrypt_executor = ThreadPoolExecutor(max_workers=2)

import bcrypt as _bcrypt
import pdfplumber
from dotenv import load_dotenv
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
from google import genai as google_genai
from google.genai import types as genai_types
from sqlalchemy import func as sqlfunc
from sqlalchemy.orm import Session
from starlette.middleware.sessions import SessionMiddleware

import models
from database import engine, get_db, run_migrations

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("gestor")

models.Base.metadata.create_all(bind=engine)
run_migrations()

DATA_DIR = Path("data") / "consultas"
DATA_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="Gestor de Prompts API")
app.add_middleware(
    SessionMiddleware,
    secret_key=os.getenv("SECRET_KEY", "change-this-in-production"),
    max_age=86400 * 7,
)
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r".*",   # LAN deployment: accept any origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin123")
OPENAI_MODELS = ["gpt-4o", "gpt-4o-mini", "o3-mini", "o1-mini", "gpt-4-turbo", "gpt-3.5-turbo"]
GEMINI_MODELS = ["gemini-2.0-flash", "gemini-1.5-pro", "gemini-1.5-flash"]
MANUAL_MODELS = ["chatgpt-manual"]
AVAILABLE_MODELS = OPENAI_MODELS + GEMINI_MODELS + MANUAL_MODELS
COMPANY_NAME = os.getenv("COMPANY_NAME", "Pagola & Madorran")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _openai():
    key = os.getenv("OPENAI_API_KEY", "")
    if not key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY no configurada en .env")
    return OpenAI(api_key=key)

def _is_gemini(model: str) -> bool:
    return model.startswith("gemini-")

def _is_admin(req: Request) -> bool:
    return req.session.get("is_admin", False)

def _me(req: Request) -> Optional[dict]:
    return req.session.get("user")

def _require(req: Request) -> dict:
    u = _me(req)
    if not u:
        raise HTTPException(status_code=401, detail="No autenticado")
    return u

def _require_admin(req: Request):
    if not _is_admin(req):
        raise HTTPException(status_code=403, detail="Acceso denegado")

def _ip(req: Request) -> str:
    fwd = req.headers.get("X-Forwarded-For")
    return fwd.split(",")[0].strip() if fwd else (req.client.host if req.client else "—")

def _safe(name: str) -> str:
    name = os.path.basename(name)
    return (re.sub(r"[^\w.\-]", "_", name) or "archivo")[:200]

def _hash_sync(pw: str) -> str:
    return _bcrypt.hashpw(pw.encode(), _bcrypt.gensalt()).decode()

def _verify_sync(pw: str, hashed: str) -> bool:
    return _bcrypt.checkpw(pw.encode(), hashed.encode())

async def _hash(pw: str) -> str:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(_bcrypt_executor, _hash_sync, pw)

async def _verify(pw: str, hashed: str) -> bool:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(_bcrypt_executor, _verify_sync, pw, hashed)

def _sections_for_user(db: Session, user_id: int) -> list:
    restricted = {r.section_id for r in db.query(models.UserSection).all()}
    allowed = {r.section_id for r in db.query(models.UserSection)
               .filter(models.UserSection.user_id == user_id).all()}
    return [s for s in db.query(models.Section)
            .filter(models.Section.deleted_at == None)
            .order_by(models.Section.sort_order).all()
            if s.id not in restricted or s.id in allowed]

def _sec_dict(s: models.Section) -> dict:
    return {
        "id": s.id, "name": s.name, "icon": s.icon, "input_type": s.input_type,
        "model": s.model, "temperature": s.temperature, "max_tokens": s.max_tokens,
        "sort_order": s.sort_order, "prompt": s.prompt,
        "description": s.description or "",
        "quick_inputs": json.loads(s.quick_inputs or "[]"),
        "created_at": s.created_at.isoformat() if s.created_at else None,
        "deleted_at": s.deleted_at.isoformat() if s.deleted_at else None,
    }

def _user_dict(u: models.User) -> dict:
    return {
        "id": u.id, "username": u.username, "display_name": u.display_name,
        "department": u.department, "is_active": u.is_active,
        "daily_limit": u.daily_limit,
        "monthly_limit": u.monthly_limit,
        "created_at": u.created_at.isoformat() if u.created_at else None,
        "deleted_at": u.deleted_at.isoformat() if u.deleted_at else None,
    }

def _query_dict(q: models.Query) -> dict:
    return {
        "id": q.id, "user_id": q.user_id, "user_display_name": q.user_display_name,
        "section_id": q.section_id, "section_name": q.section_name, "section_icon": q.section_icon,
        "client_text": q.client_text, "result": q.result,
        "has_file": q.has_file, "filename": q.filename,
        "input_tokens": q.input_tokens, "output_tokens": q.output_tokens,
        "duration_ms": q.duration_ms, "status": q.status, "error_msg": q.error_msg,
        "tags": json.loads(q.tags or "[]"),
        "is_protected": bool(q.is_protected),
        "is_favorite": bool(q.is_favorite),
        "is_comparison": bool(q.is_comparison),
        "model_b": q.model_b,
        "result_b": q.result_b,
        "created_at": q.created_at.isoformat() if q.created_at else None,
    }


# ── File extraction ───────────────────────────────────────────────────────────

async def _extract(file: UploadFile, content: bytes):
    name = (file.filename or "").lower()
    mime, _ = mimetypes.guess_type(file.filename or "")
    mime = mime or "application/octet-stream"
    if name.endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages).strip()
            return text, False, mime
        except Exception as e:
            return f"[Error PDF: {e}]", False, mime
    if name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()), False, mime
        except Exception as e:
            return f"[Error DOCX: {e}]", False, mime
    if name.endswith((".xlsx", ".xls")):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"[Hoja: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        lines.append("\t".join(cells))
            return "\n".join(lines), False, mime
        except Exception as e:
            return f"[Error Excel: {e}]", False, mime
    if name.endswith((".jpg", ".jpeg", ".png", ".webp", ".gif")):
        return base64.b64encode(content).decode(), True, mime
    try:
        return content.decode("utf-8").strip(), False, mime
    except Exception:
        return "[Archivo no legible]", False, mime


def _strip_emoji(text: str) -> str:
    """Remove emoji/non-latin-1 characters for fpdf2 Helvetica compatibility."""
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def _strip_md(text: str) -> str:
    """Remove markdown syntax for PDF plain-text output."""
    import re
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'^\s*[-*+]\s+', '• ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _build_pdf(
    title: str, section_name: str, section_icon: str,
    user_name: str, created_at: str, model: str,
    content: str, logo_path: Optional[str] = None,
) -> bytes:
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            if logo_path and Path(logo_path).exists():
                try:
                    self.image(logo_path, x=12, y=10, h=12)
                except Exception:
                    pass
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(140, 144, 164)
            self.set_xy(0, 14)
            self.cell(0, 8, COMPANY_NAME, align="R")
            self.ln(18)
            self.set_draw_color(46, 50, 72)
            self.line(12, self.get_y(), self.w - 12, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-14)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(140, 144, 164)
            self.cell(0, 8, f"Generado por Gestor IA · {COMPANY_NAME} · Pág. {self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(14, 14, 14)
    pdf.add_page()

    # Section title
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(232, 233, 240)
    pdf.multi_cell(0, 10, _strip_emoji(f"{section_icon}  {section_name}"))
    pdf.ln(2)

    # Meta line
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(140, 144, 164)
    pdf.cell(0, 6, _strip_emoji(f"{user_name}   ·   {created_at}   ·   {model}"), align="L")
    pdf.ln(10)

    # Divider
    pdf.set_draw_color(46, 50, 72)
    pdf.line(14, pdf.get_y(), pdf.w - 14, pdf.get_y())
    pdf.ln(8)

    # Content
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(232, 233, 240)
    clean = _strip_emoji(_strip_md(content))
    pdf.set_fill_color(15, 17, 23)
    pdf.rect(14, pdf.get_y() - 2, pdf.w - 28, pdf.h - pdf.get_y() - 16, "F")
    for para in clean.split("\n"):
        para = para.strip()
        if para:
            pdf.multi_cell(0, 6, para)
            pdf.ln(2)
        else:
            pdf.ln(3)

    return bytes(pdf.output())


def _build_historial_pdf(items: list, user_name: str) -> bytes:
    from fpdf import FPDF
    logo = str(Path("asserts/logo.png")) if Path("asserts/logo.png").exists() else None

    class PDF(FPDF):
        def header(self):
            if logo and Path(logo).exists():
                try:
                    self.image(logo, x=12, y=10, h=12)
                except Exception:
                    pass
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(140, 144, 164)
            self.set_xy(0, 14)
            self.cell(0, 8, COMPANY_NAME, align="R")
            self.ln(18)
            self.set_draw_color(46, 50, 72)
            self.line(12, self.get_y(), self.w - 12, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-14)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(140, 144, 164)
            self.cell(0, 8, f"Historial IA · {COMPANY_NAME} · Pág. {self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(14, 14, 14)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(232, 233, 240)
    pdf.cell(0, 10, "Historial de Consultas", ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(140, 144, 164)
    pdf.cell(0, 6, _strip_emoji(f"{user_name}  ·  {datetime.now().strftime('%d/%m/%Y %H:%M')}  ·  {len(items)} consultas"), ln=True)
    pdf.ln(6)
    pdf.set_draw_color(46, 50, 72)
    pdf.line(14, pdf.get_y(), pdf.w - 14, pdf.get_y())
    pdf.ln(6)

    for idx, q in enumerate(items):
        created = q.created_at.strftime("%d/%m/%Y %H:%M") if q.created_at else ""
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(220, 225, 240)
        pdf.multi_cell(0, 6, _strip_emoji(f"{q.section_icon or ''}  {q.section_name or '—'}"))
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(140, 144, 164)
        tokens = (q.input_tokens or 0) + (q.output_tokens or 0)
        pdf.cell(0, 5, f"{created}  ·  {'OK' if q.status == 'ok' else (q.status or '')}  ·  {tokens} tokens", ln=True)
        pdf.ln(2)
        if q.client_text:
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(160, 170, 200)
            snippet = _strip_emoji(q.client_text[:400] + ("..." if len(q.client_text) > 400 else ""))
            pdf.multi_cell(0, 5, snippet)
            pdf.ln(1)
        if q.result and q.status == "ok":
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(200, 210, 230)
            answer = _strip_emoji(_strip_md(q.result[:1000] + ("..." if len(q.result) > 1000 else "")))
            for para in answer.split("\n"):
                para = para.strip()
                if para:
                    pdf.multi_cell(0, 5, para)
                else:
                    pdf.ln(2)
        pdf.ln(4)
        if idx < len(items) - 1:
            pdf.set_draw_color(35, 40, 60)
            pdf.line(14, pdf.get_y(), pdf.w - 14, pdf.get_y())
            pdf.ln(4)

    return bytes(pdf.output())


def _save_disk(qid: int, text: str, result: str, content: Optional[bytes], fname: Optional[str]) -> str:
    folder = DATA_DIR / str(qid)
    folder.mkdir(parents=True, exist_ok=True)
    if text and text.strip():
        (folder / "pregunta.txt").write_text(text, encoding="utf-8")
    if result:
        (folder / "respuesta.md").write_text(result, encoding="utf-8")
    if content and fname:
        d = folder / "archivos"
        d.mkdir(exist_ok=True)
        (d / _safe(fname)).write_bytes(content)
    return str(folder)


# ── Seed sections ─────────────────────────────────────────────────────────────

SEED_SECTIONS = [
    {"name": "Analizador de Pólizas", "icon": "📋", "input_type": "both", "model": "gpt-4o",
     "temperature": 0.3, "max_tokens": 3000, "sort_order": 1,
     "prompt": 'Eres un experto en seguros con 20 años de experiencia.\n\nCommentario del agente: "{TEXTO_CLIENTE}"\n\nDocumento:\n{TEXTO_DEL_ARCHIVO}\n\n## 📋 Resumen\n## ✅ Coberturas\n## ❌ Exclusiones\n## ⚠️ Franquicias y carencias\n## 💡 Puntos de atención'},
    {"name": "Redactor de Comunicaciones", "icon": "✉️", "input_type": "text", "model": "gpt-4o",
     "temperature": 0.7, "max_tokens": 2000, "sort_order": 2,
     "prompt": 'Eres un redactor profesional de seguros.\n\nInstrucciones: "{TEXTO_CLIENTE}"\n\nDocumento de referencia:\n{TEXTO_DEL_ARCHIVO}\n\n## ✉️ Comunicación principal\n## 📝 Versión alternativa\n## 💬 Puntos clave'},
    {"name": "Gestor de Siniestros", "icon": "🚨", "input_type": "both", "model": "gpt-4o",
     "temperature": 0.3, "max_tokens": 3000, "sort_order": 3,
     "prompt": 'Eres un especialista en gestión de siniestros.\n\nNotas: "{TEXTO_CLIENTE}"\n\nDocumentación:\n{TEXTO_DEL_ARCHIVO}\n\n## 🚨 Resumen\n## 📊 Análisis de cobertura\n## 📋 Documentación\n## 🔍 Puntos de atención\n## ✅ Próximos pasos\n## 💰 Estimación provisional'},
    {"name": "Comparador de Coberturas", "icon": "⚖️", "input_type": "both", "model": "gpt-4o",
     "temperature": 0.4, "max_tokens": 3000, "sort_order": 4,
     "prompt": 'Eres un analista de seguros experto en comparativas.\n\nIndicaciones: "{TEXTO_CLIENTE}"\n\nDocumentación:\n{TEXTO_DEL_ARCHIVO}\n\n## 📊 Tabla comparativa\n## ✅ Ventajas e inconvenientes\n## 🏆 Recomendación\n## ⚠️ Aspectos a negociar'},
]

def _seed(db: Session):
    if db.query(models.Section).count() == 0:
        for d in SEED_SECTIONS:
            db.add(models.Section(**d))
        db.commit()


# ══════════════════════════════════════════════════════════════════════════════
# USER AUTH
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/me")
async def api_me(req: Request):
    u = _me(req)
    if not u:
        raise HTTPException(status_code=401, detail="No autenticado")
    return u

@app.post("/api/login")
async def api_login(req: Request, body: Dict[str, Any], db: Session = Depends(get_db)):
    username = body.get("username", "").strip().lower()
    password = body.get("password", "")
    user = db.query(models.User).filter(
        models.User.username == username, models.User.is_active == True
    ).first()
    if not user or not await _verify(password, user.password_hash):
        log.warning("Failed login: '%s' from %s", username, _ip(req))
        raise HTTPException(status_code=401, detail="Usuario o contraseña incorrectos")
    req.session["user"] = {
        "id": user.id, "username": user.username,
        "display_name": user.display_name, "department": user.department,
    }
    log.info("Login: '%s' from %s", username, _ip(req))
    return {"ok": True, "user": _user_dict(user)}

@app.post("/api/logout")
async def api_logout(req: Request):
    req.session.pop("user", None)
    return {"ok": True}

@app.get("/health")
async def health():
    return {"ok": True}

@app.get("/api/config")
async def api_config():
    return {"company_name": COMPANY_NAME}

@app.get("/api/sections")
async def api_sections(req: Request, db: Session = Depends(get_db)):
    u = _require(req)
    return [_sec_dict(s) for s in _sections_for_user(db, u["id"])]


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN AUTH
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/models")
async def api_admin_models(req: Request):
    _require_admin(req)
    return {
        "openai": OPENAI_MODELS,
        "gemini": GEMINI_MODELS,
        "manual": MANUAL_MODELS,
        "all": AVAILABLE_MODELS,
    }

@app.get("/api/admin/me")
async def api_admin_me(req: Request):
    if not _is_admin(req):
        raise HTTPException(status_code=403, detail="Acceso denegado")
    return {"is_admin": True}

@app.post("/api/admin/login")
async def api_admin_login(req: Request, body: Dict[str, Any]):
    if body.get("password") != ADMIN_PASSWORD:
        log.warning("Failed admin login from %s", _ip(req))
        raise HTTPException(status_code=401, detail="Contraseña incorrecta")
    req.session["is_admin"] = True
    log.info("Admin login from %s", _ip(req))
    return {"ok": True}

@app.post("/api/admin/logout")
async def api_admin_logout(req: Request):
    req.session.pop("is_admin", None)
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — STATS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/stats")
async def api_admin_stats(req: Request, db: Session = Depends(get_db)):
    _require_admin(req)
    _seed(db)
    total_q = db.query(models.Query).count()
    ok_q = db.query(models.Query).filter(models.Query.status == "ok").count()
    err_q = total_q - ok_q
    avg_dur = int(db.query(sqlfunc.avg(models.Query.duration_ms)).scalar() or 0)
    total_tok = db.query(
        sqlfunc.sum(models.Query.input_tokens + models.Query.output_tokens)
    ).scalar() or 0

    # Per-section query counts
    from sqlalchemy import case
    sec_stats = db.query(
        models.Section.id, models.Section.name, models.Section.icon,
        sqlfunc.count(models.Query.id).label("count"),
        sqlfunc.sum(case((models.Query.status == "ok", 1), else_=0)).label("ok"),
    ).outerjoin(models.Query, models.Query.section_id == models.Section.id)\
     .group_by(models.Section.id).all()

    return {
        "sections": db.query(models.Section).count(),
        "users": db.query(models.User).count(),
        "active_users": db.query(models.User).filter(models.User.is_active == True).count(),
        "total_queries": total_q,
        "ok_queries": ok_q,
        "err_queries": err_q,
        "avg_duration_ms": avg_dur,
        "total_tokens": int(total_tok),
        "per_section": [
            {"id": r.id, "name": r.name, "icon": r.icon,
             "count": r.count or 0, "ok": r.ok or 0}
            for r in sec_stats
        ],
    }


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — SECTIONS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/sections")
async def api_admin_sections(req: Request, include_deleted: bool = False, db: Session = Depends(get_db)):
    _require_admin(req)
    _seed(db)
    q = db.query(models.Section).order_by(models.Section.sort_order)
    if not include_deleted:
        q = q.filter(models.Section.deleted_at == None)
    secs = q.all()
    result = []
    for s in secs:
        d = _sec_dict(s)
        assigned = [
            {"user_id": r.user_id, "daily_limit": r.daily_limit}
            for r in db.query(models.UserSection).filter(models.UserSection.section_id == s.id).all()
        ]
        d["allowed_users"] = assigned
        result.append(d)
    return result

@app.post("/api/admin/sections")
async def api_admin_create_section(req: Request, body: Dict[str, Any], db: Session = Depends(get_db)):
    _require_admin(req)
    s = models.Section(
        name=body["name"], prompt=body["prompt"],
        input_type=body.get("input_type", "both"),
        icon=body.get("icon", "⚡"),
        model=body.get("model", "gpt-4o") if body.get("model") in AVAILABLE_MODELS else "gpt-4o",
        description=body.get("description", ""),
        quick_inputs=json.dumps(body.get("quick_inputs", [])),
        temperature=max(0.0, min(2.0, float(body.get("temperature", 0.7)))),
        max_tokens=max(100, min(8000, int(body.get("max_tokens", 2000)))),
        sort_order=int(body.get("sort_order", 0)),
    )
    db.add(s)
    db.commit()
    db.refresh(s)
    _update_permissions(db, s.id, body.get("allowed_users", []))
    _audit(db, req, "section.create", "section", s.id, s.name)
    db.commit()
    log.info("Section created: '%s'", s.name)
    return _sec_dict(s)

@app.get("/api/admin/sections/{sid}")
async def api_admin_get_section(req: Request, sid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    s = db.query(models.Section).filter(models.Section.id == sid).first()
    if not s:
        raise HTTPException(status_code=404)
    d = _sec_dict(s)
    d["allowed_users"] = [
        {"user_id": r.user_id, "daily_limit": r.daily_limit}
        for r in db.query(models.UserSection).filter(models.UserSection.section_id == sid).all()
    ]
    return d

@app.put("/api/admin/sections/{sid}")
async def api_admin_update_section(req: Request, sid: int, body: Dict[str, Any], db: Session = Depends(get_db)):
    _require_admin(req)
    s = db.query(models.Section).filter(models.Section.id == sid).first()
    if not s:
        raise HTTPException(status_code=404)
    # Save current version before any changes
    db.add(models.SectionVersion(
        section_id=s.id, section_name=s.name,
        prompt=s.prompt, model=s.model,
        temperature=s.temperature, max_tokens=s.max_tokens,
        changed_by_ip=_ip(req),
    ))
    s.name = body.get("name", s.name)
    s.prompt = body.get("prompt", s.prompt)
    s.input_type = body.get("input_type", s.input_type)
    s.icon = body.get("icon", s.icon) or "⚡"
    if body.get("model") in AVAILABLE_MODELS:
        s.model = body["model"]
    s.temperature = max(0.0, min(2.0, float(body.get("temperature", s.temperature))))
    s.max_tokens = max(100, min(8000, int(body.get("max_tokens", s.max_tokens))))
    s.sort_order = int(body.get("sort_order", s.sort_order))
    if "description" in body:
        s.description = body["description"]
    if "quick_inputs" in body:
        s.quick_inputs = json.dumps(body["quick_inputs"] if isinstance(body["quick_inputs"], list) else [])
    _update_permissions(db, sid, body.get("allowed_users", None))
    _audit(db, req, "section.update", "section", s.id, s.name)
    db.commit()
    log.info("Section updated: '%s'", s.name)
    return _sec_dict(s)

@app.delete("/api/admin/sections/{sid}")
async def api_admin_delete_section(req: Request, sid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    s = db.query(models.Section).filter(models.Section.id == sid).first()
    if s:
        _audit(db, req, "section.delete", "section", s.id, s.name)
        s.deleted_at = datetime.now()
        db.commit()
        log.info("Section soft-deleted: '%s'", s.name)
    return {"ok": True}

@app.post("/api/admin/sections/{sid}/restore")
async def api_admin_restore_section(req: Request, sid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    s = db.query(models.Section).filter(models.Section.id == sid).first()
    if not s:
        raise HTTPException(status_code=404)
    s.deleted_at = None
    _audit(db, req, "section.restore", "section", s.id, s.name)
    db.commit()
    return _sec_dict(s)

@app.post("/api/admin/sections/{sid}/duplicate")
async def api_admin_dup_section(req: Request, sid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    s = db.query(models.Section).filter(models.Section.id == sid).first()
    if not s:
        raise HTTPException(status_code=404)
    new_s = models.Section(
        name=f"{s.name} (copia)", prompt=s.prompt, input_type=s.input_type,
        icon=s.icon, model=s.model, temperature=s.temperature,
        max_tokens=s.max_tokens, sort_order=s.sort_order + 1,
    )
    db.add(new_s)
    db.commit()
    db.refresh(new_s)
    return _sec_dict(new_s)

@app.post("/api/admin/sections/reorder")
async def api_admin_reorder(req: Request, body: List[Dict], db: Session = Depends(get_db)):
    _require_admin(req)
    for item in body:
        s = db.query(models.Section).filter(models.Section.id == item["id"]).first()
        if s:
            s.sort_order = item["sort_order"]
    db.commit()
    return {"ok": True}

def _update_permissions(db: Session, sid: int, allowed_users):
    if allowed_users is None:
        return
    db.query(models.UserSection).filter(models.UserSection.section_id == sid).delete()
    for item in allowed_users:
        if isinstance(item, dict):
            uid = int(item["user_id"])
            dl = item.get("daily_limit")
            dl = int(dl) if dl else None
        else:
            uid = int(item)
            dl = None
        db.add(models.UserSection(user_id=uid, section_id=sid, daily_limit=dl))
    db.commit()


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — USERS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/users")
async def api_admin_users(req: Request, include_deleted: bool = False, db: Session = Depends(get_db)):
    _require_admin(req)
    q = db.query(models.User)
    if not include_deleted:
        q = q.filter(models.User.deleted_at == None)
    users = q.all()
    result = []
    for u in users:
        d = _user_dict(u)
        d["query_count"] = db.query(models.Query).filter(models.Query.user_id == u.id).count()
        d["section_ids"] = [r.section_id for r in db.query(models.UserSection)
                            .filter(models.UserSection.user_id == u.id).all()]
        result.append(d)
    return result

@app.post("/api/admin/users")
async def api_admin_create_user(req: Request, body: Dict[str, Any], db: Session = Depends(get_db)):
    _require_admin(req)
    uname = body.get("username", "").strip().lower()
    if not uname:
        raise HTTPException(status_code=400, detail="Nombre de usuario requerido")
    if db.query(models.User).filter(models.User.username == uname).first():
        raise HTTPException(status_code=409, detail=f"Usuario '{uname}' ya existe")
    u = models.User(
        username=uname,
        display_name=body.get("display_name", uname).strip(),
        department=body.get("department", "").strip(),
        password_hash=await _hash(body.get("password", "changeme")),
    )
    db.add(u)
    db.commit()
    db.refresh(u)
    _audit(db, req, "user.create", "user", u.id, u.display_name)
    db.commit()
    log.info("User created: '%s'", uname)
    return _user_dict(u)

@app.put("/api/admin/users/{uid}")
async def api_admin_update_user(req: Request, uid: int, body: Dict[str, Any], db: Session = Depends(get_db)):
    _require_admin(req)
    u = db.query(models.User).filter(models.User.id == uid).first()
    if not u:
        raise HTTPException(status_code=404)
    u.display_name = body.get("display_name", u.display_name).strip()
    u.department = body.get("department", u.department).strip()
    if body.get("password"):
        u.password_hash = await _hash(body["password"])
    if "daily_limit" in body:
        u.daily_limit = int(body["daily_limit"]) if body["daily_limit"] else None
    if "monthly_limit" in body:
        u.monthly_limit = int(body["monthly_limit"]) if body["monthly_limit"] else None
    _audit(db, req, "user.update", "user", u.id, u.display_name)
    db.commit()
    return _user_dict(u)

@app.post("/api/admin/users/{uid}/toggle")
async def api_admin_toggle_user(req: Request, uid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    u = db.query(models.User).filter(models.User.id == uid).first()
    if not u:
        raise HTTPException(status_code=404)
    u.is_active = not u.is_active
    _audit(db, req, "user.toggle", "user", u.id, u.display_name,
           {"is_active": u.is_active})
    db.commit()
    return _user_dict(u)

@app.delete("/api/admin/users/{uid}")
async def api_admin_delete_user(req: Request, uid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    u = db.query(models.User).filter(models.User.id == uid).first()
    if u:
        _audit(db, req, "user.delete", "user", u.id, u.display_name)
        u.deleted_at = datetime.now()
        u.is_active = False
        db.commit()
        log.info("User soft-deleted: '%s'", u.username)
    return {"ok": True}

@app.post("/api/admin/users/{uid}/restore")
async def api_admin_restore_user(req: Request, uid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    u = db.query(models.User).filter(models.User.id == uid).first()
    if not u:
        raise HTTPException(status_code=404)
    u.deleted_at = None
    u.is_active = True
    _audit(db, req, "user.restore", "user", u.id, u.display_name)
    db.commit()
    return _user_dict(u)

@app.put("/api/admin/users/{uid}/sections")
async def api_admin_user_set_sections(req: Request, uid: int, body: Dict[str, Any], db: Session = Depends(get_db)):
    _require_admin(req)
    u = db.query(models.User).filter(models.User.id == uid).first()
    if not u:
        raise HTTPException(status_code=404)
    db.query(models.UserSection).filter(models.UserSection.user_id == uid).delete()
    for entry in body.get("sections", []):
        sid = entry.get("section_id")
        dl = entry.get("daily_limit")
        if sid:
            db.add(models.UserSection(user_id=uid, section_id=sid, daily_limit=dl))
    _audit(db, req, "user.sections.update", "user", uid, u.display_name,
           {"section_ids": [e.get("section_id") for e in body.get("sections", [])]})
    db.commit()
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — QUERIES (all users)
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/queries")
async def api_admin_queries(
    req: Request, page: int = 1,
    section_id: Optional[int] = None,
    user_id: Optional[int] = None,
    status: Optional[str] = None,
    q: Optional[str] = None,
    db: Session = Depends(get_db),
):
    _require_admin(req)
    from sqlalchemy import or_
    per_page = 40
    query = db.query(models.Query).order_by(models.Query.created_at.desc())
    if section_id:
        query = query.filter(models.Query.section_id == section_id)
    if user_id:
        query = query.filter(models.Query.user_id == user_id)
    if status:
        query = query.filter(models.Query.status == status)
    if q and q.strip():
        term = f"%{q.strip()}%"
        query = query.filter(or_(
            models.Query.client_text.ilike(term),
            models.Query.section_name.ilike(term),
            models.Query.result.ilike(term),
            models.Query.user_display_name.ilike(term),
        ))
    total = query.count()
    items = query.offset((page - 1) * per_page).limit(per_page).all()
    return {
        "items": [_query_dict(i) for i in items],
        "total": total,
        "pages": max(1, (total + per_page - 1) // per_page),
        "page": page,
    }

@app.get("/api/admin/queries/export")
async def api_admin_queries_export(
    req: Request,
    format: str = "csv",
    section_id: Optional[int] = None,
    user_id: Optional[int] = None,
    status: Optional[str] = None,
    db: Session = Depends(get_db),
):
    _require_admin(req)
    from sqlalchemy import or_
    query = db.query(models.Query).order_by(models.Query.created_at.desc())
    if section_id:
        query = query.filter(models.Query.section_id == section_id)
    if user_id:
        query = query.filter(models.Query.user_id == user_id)
    if status:
        query = query.filter(models.Query.status == status)
    items = query.limit(10000).all()

    if format == "pdf":
        pdf_bytes = _build_historial_pdf(items, "Admin")
        fname = f"consultas_admin_{datetime.now().strftime('%Y%m%d')}.pdf"
        return StreamingResponse(iter([pdf_bytes]), media_type="application/pdf",
                                 headers={"Content-Disposition": f"attachment; filename={fname}"})

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["id", "fecha", "usuario", "seccion", "consulta", "resultado",
                     "tokens_entrada", "tokens_salida", "estado", "duracion_ms", "tiene_archivo"])
    for q in items:
        writer.writerow([
            q.id,
            q.created_at.isoformat() if q.created_at else "",
            q.user_display_name or "",
            q.section_name or "",
            (q.client_text or "").replace("\n", " "),
            (q.result or "").replace("\n", " ")[:500],
            q.input_tokens or 0, q.output_tokens or 0,
            q.status or "", q.duration_ms or 0,
            "sí" if q.has_file else "no",
        ])
    output.seek(0)
    fname = f"consultas_admin_{datetime.now().strftime('%Y%m%d')}.csv"
    return StreamingResponse(io.BytesIO(output.getvalue().encode("utf-8-sig")),
                             media_type="text/csv",
                             headers={"Content-Disposition": f"attachment; filename={fname}"})


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — LOGS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/logs")
async def api_admin_logs(
    req: Request, page: int = 1,
    user: Optional[str] = None,
    section: Optional[str] = None,
    status: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    db: Session = Depends(get_db),
):
    _require_admin(req)
    per_page = 50
    query = db.query(models.Log)
    if user:
        query = query.filter(models.Log.user_display_name.ilike(f"%{user}%"))
    if section:
        query = query.filter(models.Log.section_name.ilike(f"%{section}%"))
    if status:
        query = query.filter(models.Log.status == status)
    if date_from:
        query = query.filter(models.Log.created_at >= datetime.strptime(date_from, "%Y-%m-%d"))
    if date_to:
        query = query.filter(models.Log.created_at < datetime.strptime(date_to, "%Y-%m-%d") + timedelta(days=1))
    total = query.count()
    items = query.order_by(models.Log.created_at.desc()).offset((page - 1) * per_page).limit(per_page).all()
    ok = db.query(models.Log).filter(models.Log.status == "ok").count()
    err_total = db.query(models.Log).count()
    avg = int(db.query(sqlfunc.avg(models.Log.duration_ms)).scalar() or 0)
    return {
        "items": [{
            "id": l.id, "user_display_name": l.user_display_name,
            "section_name": l.section_name, "client_ip": l.client_ip,
            "has_file": l.has_file, "filename": l.filename,
            "input_tokens": l.input_tokens, "output_tokens": l.output_tokens,
            "status": l.status, "error_msg": l.error_msg,
            "duration_ms": l.duration_ms,
            "created_at": l.created_at.isoformat() if l.created_at else None,
        } for l in items],
        "total": total, "pages": max(1, (total + per_page - 1) // per_page),
        "page": page, "ok_count": ok, "err_count": err_total - ok, "avg_duration_ms": avg,
    }

@app.delete("/api/admin/logs")
async def api_admin_clear_logs(req: Request, db: Session = Depends(get_db)):
    _require_admin(req)
    n = db.query(models.Log).delete()
    db.commit()
    log.info("Logs cleared: %d rows", n)
    return {"ok": True, "deleted": n}

@app.get("/api/admin/logs/export")
async def api_admin_export_logs(req: Request, db: Session = Depends(get_db)):
    _require_admin(req)
    items = db.query(models.Log).order_by(models.Log.created_at.desc()).all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["id", "fecha", "usuario", "seccion", "ip", "archivo",
                     "tokens_entrada", "tokens_salida", "estado", "duracion_ms", "error"])
    for l in items:
        writer.writerow([
            l.id, l.created_at.isoformat() if l.created_at else "",
            l.user_display_name or "", l.section_name or "", l.client_ip or "",
            l.filename or "", l.input_tokens, l.output_tokens,
            l.status, l.duration_ms, l.error_msg or "",
        ])
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=logs.csv"},
    )


# ══════════════════════════════════════════════════════════════════════════════
# HISTORIAL (user)
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/historial/export")
async def api_historial_export(req: Request, format: str = "csv", db: Session = Depends(get_db)):
    u = _require(req)
    items = (db.query(models.Query)
             .filter(models.Query.user_id == u["id"])
             .order_by(models.Query.created_at.desc())
             .all())

    if format == "pdf":
        pdf_bytes = _build_historial_pdf(items, u.get("display_name", u["username"]))
        filename = f"historial_{u['username']}_{datetime.now().strftime('%Y%m%d')}.pdf"
        return StreamingResponse(
            iter([pdf_bytes]),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["id", "fecha", "seccion", "consulta", "resultado", "tokens_entrada",
                     "tokens_salida", "estado", "duracion_ms", "tiene_archivo", "tags"])
    for q in items:
        writer.writerow([
            q.id,
            q.created_at.isoformat() if q.created_at else "",
            q.section_name or "",
            (q.client_text or "").replace("\n", " "),
            (q.result or "").replace("\n", " ")[:500],
            q.input_tokens or 0,
            q.output_tokens or 0,
            q.status or "",
            q.duration_ms or 0,
            "sí" if q.has_file else "no",
            q.tags or "[]",
        ])
    output.seek(0)
    filename = f"historial_{u['username']}_{datetime.now().strftime('%Y%m%d')}.csv"
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@app.get("/api/historial/{qid}")
async def api_historial_detail(req: Request, qid: int, db: Session = Depends(get_db)):
    u = _require(req)
    item = db.query(models.Query).filter(
        models.Query.id == qid, models.Query.user_id == u["id"]
    ).first()
    if not item:
        raise HTTPException(status_code=404)
    files = []
    if item.folder_path:
        d = Path(item.folder_path) / "archivos"
        if d.exists():
            files = [f.name for f in d.iterdir() if f.is_file()]
    return {**_query_dict(item), "files": files}

@app.get("/api/historial/{qid}/pdf")
async def api_export_pdf(req: Request, qid: int, db: Session = Depends(get_db)):
    u = _require(req)
    item = db.query(models.Query).filter(
        models.Query.id == qid, models.Query.user_id == u["id"]
    ).first()
    if not item:
        raise HTTPException(status_code=404)
    logo = str(Path("asserts/logo.png")) if Path("asserts/logo.png").exists() else None
    created = item.created_at.strftime("%d/%m/%Y %H:%M") if item.created_at else ""
    sec = db.query(models.Section).filter(models.Section.id == item.section_id).first()
    pdf_bytes = _build_pdf(
        title=item.section_name or "Consulta",
        section_name=item.section_name or "Consulta",
        section_icon=item.section_icon or "📄",
        user_name=item.user_display_name or u["display_name"],
        created_at=created,
        model=sec.model if sec else "",
        content=item.result or "",
        logo_path=logo,
    )
    safe_name = re.sub(r"[^\w\-]", "_", item.section_name or "consulta")
    return StreamingResponse(
        iter([pdf_bytes]),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={safe_name}_{qid}.pdf"},
    )

@app.get("/api/historial/{qid}/archivo/{filename}")
async def api_download_file(req: Request, qid: int, filename: str, db: Session = Depends(get_db)):
    u = _require(req)
    item = db.query(models.Query).filter(
        models.Query.id == qid, models.Query.user_id == u["id"]
    ).first()
    if not item or not item.folder_path:
        raise HTTPException(status_code=404)
    path = Path(item.folder_path) / "archivos" / _safe(filename)
    if not path.exists():
        raise HTTPException(status_code=404)
    return FileResponse(path, filename=_safe(filename))


# ══════════════════════════════════════════════════════════════════════════════
# SHARED PROCESS HELPERS
# ══════════════════════════════════════════════════════════════════════════════

async def _prepare_request(section, client_text: str, files: list, user: dict):
    """Extract files, compile prompt with all variables, build messages."""
    all_contents: List[bytes] = []
    all_names: List[str] = []
    text_parts: List[str] = []
    is_image = False
    image_mime = "image/jpeg"
    image_b64 = ""

    for f in files:
        if not f or not f.filename:
            continue
        content = await f.read()
        all_contents.append(content)
        all_names.append(f.filename)
        ft, img, mime = await _extract(f, content)
        if img and not is_image:
            is_image = True
            image_mime = mime
            image_b64 = ft
        elif not img:
            label = f"[{f.filename}]" if len(files) > 1 else ""
            text_parts.append((label + "\n" + ft).strip())

    combined = "\n\n---\n\n".join(text_parts)
    file_content = all_contents[0] if all_contents else None
    filename = all_names[0] if all_names else None

    prompt = section.prompt
    # Support both legacy {VAR} format and modern {{var}} format
    file_text_val = "[imagen adjunta]" if is_image else combined
    fname_val = ", ".join(all_names) if all_names else ""
    now_str = datetime.now().strftime("%d/%m/%Y")
    # Legacy format
    prompt = prompt.replace("{TEXTO_CLIENTE}", client_text or "")
    prompt = prompt.replace("{TEXTO_DEL_ARCHIVO}", file_text_val)
    prompt = prompt.replace("{FECHA}", now_str)
    prompt = prompt.replace("{USUARIO}", user.get("display_name", ""))
    prompt = prompt.replace("{DEPARTAMENTO}", user.get("department", ""))
    # Modern format
    prompt = prompt.replace("{{text}}", client_text or "")
    prompt = prompt.replace("{{file_content}}", file_text_val)
    prompt = prompt.replace("{{filename}}", fname_val)
    prompt = prompt.replace("{{user_name}}", user.get("display_name", ""))
    prompt = prompt.replace("{{department}}", user.get("department", ""))
    prompt = prompt.replace("{{date}}", now_str)

    messages = [{
        "role": "user",
        "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:{image_mime};base64,{image_b64}"}},
        ] if is_image else prompt,
    }]

    return prompt, messages, file_content, filename, all_names, is_image, image_mime, image_b64


# ══════════════════════════════════════════════════════════════════════════════
# PROCESS (standard)
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/process/{sid}")
async def api_process(
    req: Request, sid: int,
    client_text: Optional[str] = Form(None),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
):
    u = _require(req)
    _check_limits(db, u["id"], sid)
    section = db.query(models.Section).filter(models.Section.id == sid).first()
    if not section:
        raise HTTPException(status_code=404)

    t0 = int(time.time() * 1000)
    valid_files = [f for f in files if f and f.filename]

    # Cache check — only for text-only queries (no files)
    if not valid_files and client_text and client_text.strip() and section.model != "chatgpt-manual":
        ckey = _cache_key(u["id"], sid, client_text.strip())
        cached = _cache_get(ckey)
        if cached:
            log.info("Cache hit '%s' | %s", section.name, u["username"])
            return JSONResponse({**cached, "from_cache": True})
    prompt, messages, file_content, filename, all_names, is_image, image_mime, image_b64 = \
        await _prepare_request(section, client_text or "", valid_files, u)
    file_text = image_b64 if is_image else ""

    log_e = models.Log(
        user_display_name=u["display_name"],
        section_id=section.id, section_name=section.name,
        client_ip=_ip(req), has_file=bool(filename), filename=filename, status="error",
    )

    try:
        if section.model == "chatgpt-manual":
            # Don't call any API — return the compiled prompt for the user to paste into ChatGPT
            dur = int(time.time() * 1000) - t0
            log_e.status = "ok"
            log_e.input_tokens = 0
            log_e.output_tokens = 0
            log_e.duration_ms = dur
            qrow = models.Query(
                user_id=u["id"], user_display_name=u["display_name"],
                section_id=section.id, section_name=section.name, section_icon=section.icon,
                client_text=client_text or "", result=prompt,
                has_file=bool(filename), filename=filename,
                input_tokens=0, output_tokens=0,
                duration_ms=dur, status="ok",
            )
            db.add(qrow)
            db.flush()
            qrow.folder_path = _save_disk(qrow.id, client_text or "", prompt, file_content, filename)
            db.add(log_e)
            db.commit()
            return JSONResponse({
                "result": prompt,
                "query_id": qrow.id,
                "mode": "manual",
                "tokens": {"input": 0, "output": 0, "total": 0},
                "duration_ms": dur,
                "model": "chatgpt-manual",
            })

        if _is_gemini(section.model):
            gemini_key = os.getenv("GEMINI_API_KEY", "")
            if not gemini_key:
                raise HTTPException(status_code=500, detail="GEMINI_API_KEY no configurada en .env")
            gclient = google_genai.Client(api_key=gemini_key)
            gen_cfg = genai_types.GenerateContentConfig(
                temperature=section.temperature,
                max_output_tokens=section.max_tokens,
            )
            if is_image:
                import PIL.Image
                img = PIL.Image.open(io.BytesIO(base64.b64decode(image_b64)))
                gresp = gclient.models.generate_content(
                    model=section.model, contents=[prompt, img], config=gen_cfg
                )
            else:
                gresp = gclient.models.generate_content(
                    model=section.model, contents=prompt, config=gen_cfg
                )
            result = gresp.text
            in_tok = getattr(gresp.usage_metadata, "prompt_token_count", 0) or 0
            out_tok = getattr(gresp.usage_metadata, "candidates_token_count", 0) or 0
        else:
            oai = _openai()
            resp = oai.chat.completions.create(
                model=section.model, messages=messages,
                temperature=section.temperature, max_tokens=section.max_tokens,
            )
            result = resp.choices[0].message.content
            usage = resp.usage
            in_tok = usage.prompt_tokens if usage else 0
            out_tok = usage.completion_tokens if usage else 0

        dur = int(time.time() * 1000) - t0
        log_e.status = "ok"
        log_e.input_tokens = in_tok
        log_e.output_tokens = out_tok
        log_e.duration_ms = dur

        qrow = models.Query(
            user_id=u["id"], user_display_name=u["display_name"],
            section_id=section.id, section_name=section.name, section_icon=section.icon,
            client_text=client_text or "", result=result,
            has_file=bool(filename), filename=filename,
            input_tokens=in_tok, output_tokens=out_tok,
            duration_ms=dur, status="ok",
        )
        db.add(qrow)
        db.flush()
        qrow.folder_path = _save_disk(qrow.id, client_text or "", result, file_content, filename)
        db.add(log_e)
        db.commit()

        log.info("OK '%s' | %s | %d+%d tok | %dms",
                 section.name, u["username"], in_tok, out_tok, dur)

        payload = {
            "result": result, "query_id": qrow.id,
            "tokens": {"input": in_tok, "output": out_tok, "total": in_tok + out_tok},
            "duration_ms": dur, "model": section.model,
        }
        # Store in cache (text-only queries only)
        if not valid_files and client_text and client_text.strip():
            ckey = _cache_key(u["id"], sid, client_text.strip())
            _cache_set(ckey, payload)
        return JSONResponse(payload)

    except HTTPException:
        raise
    except Exception as e:
        log_e.error_msg = str(e)
        log_e.duration_ms = int(time.time() * 1000) - t0
        qrow = models.Query(
            user_id=u["id"], user_display_name=u["display_name"],
            section_id=section.id, section_name=section.name, section_icon=section.icon,
            client_text=client_text or "", result="",
            has_file=bool(filename), filename=filename,
            status="error", error_msg=str(e), duration_ms=log_e.duration_ms,
        )
        db.add(qrow)
        db.add(log_e)
        db.commit()
        log.error("Error '%s': %s", section.name, e)
        return JSONResponse({"error": str(e)}, status_code=500)


# ══════════════════════════════════════════════════════════════════════════════
# PROCESS (streaming SSE)
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/process/{sid}/stream")
async def api_process_stream(
    req: Request, sid: int,
    client_text: Optional[str] = Form(None),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
):
    u = _require(req)
    _check_limits(db, u["id"], sid)
    section = db.query(models.Section).filter(models.Section.id == sid).first()
    if not section:
        raise HTTPException(status_code=404)

    t0 = int(time.time() * 1000)
    valid_files = [f for f in files if f and f.filename]
    prompt, messages, file_content, filename, all_names, is_image, image_mime, image_b64 = \
        await _prepare_request(section, client_text or "", valid_files, u)

    user_id = u["id"]
    user_display = u["display_name"]
    section_id, section_name, section_icon = section.id, section.name, section.icon
    model_name = section.model

    async def generate():
        from database import SessionLocal as _SL
        db2 = _SL()
        full_text = ""
        in_tok = out_tok = 0
        error_msg = None
        try:
            if model_name == "chatgpt-manual":
                yield f"data: {json.dumps({'text': prompt})}\n\n"
                full_text = prompt

            elif _is_gemini(model_name):
                gkey = os.getenv("GEMINI_API_KEY", "")
                if not gkey:
                    raise ValueError("GEMINI_API_KEY no configurada en .env")
                gclient = google_genai.Client(api_key=gkey)
                cfg = genai_types.GenerateContentConfig(
                    temperature=section.temperature, max_output_tokens=section.max_tokens
                )
                if is_image:
                    import PIL.Image
                    img_obj = PIL.Image.open(io.BytesIO(base64.b64decode(image_b64)))
                    stream = gclient.models.generate_content_stream(
                        model=model_name, contents=[prompt, img_obj], config=cfg
                    )
                else:
                    stream = gclient.models.generate_content_stream(
                        model=model_name, contents=prompt, config=cfg
                    )
                for chunk in stream:
                    txt = chunk.text or ""
                    if txt:
                        full_text += txt
                        yield f"data: {json.dumps({'text': txt})}\n\n"
                    if hasattr(chunk, "usage_metadata") and chunk.usage_metadata:
                        in_tok = getattr(chunk.usage_metadata, "prompt_token_count", 0) or 0
                        out_tok = getattr(chunk.usage_metadata, "candidates_token_count", 0) or 0

            else:
                oai = _openai()
                stream = oai.chat.completions.create(
                    model=model_name, messages=messages,
                    temperature=section.temperature, max_tokens=section.max_tokens,
                    stream=True, stream_options={"include_usage": True},
                )
                for chunk in stream:
                    delta = (chunk.choices[0].delta.content or "") if chunk.choices else ""
                    if delta:
                        full_text += delta
                        yield f"data: {json.dumps({'text': delta})}\n\n"
                    if chunk.usage:
                        in_tok = chunk.usage.prompt_tokens or 0
                        out_tok = chunk.usage.completion_tokens or 0

            dur = int(time.time() * 1000) - t0
            log_e = models.Log(
                user_display_name=user_display, section_id=section_id,
                section_name=section_name, client_ip="stream",
                has_file=bool(filename), filename=filename,
                input_tokens=in_tok, output_tokens=out_tok,
                status="ok", duration_ms=dur,
            )
            qrow = models.Query(
                user_id=user_id, user_display_name=user_display,
                section_id=section_id, section_name=section_name, section_icon=section_icon,
                client_text=client_text or "", result=full_text,
                has_file=bool(filename), filename=filename,
                input_tokens=in_tok, output_tokens=out_tok, duration_ms=dur, status="ok",
            )
            db2.add(qrow)
            db2.flush()
            qrow.folder_path = _save_disk(qrow.id, client_text or "", full_text, file_content, filename)
            db2.add(log_e)
            db2.commit()
            mode = "manual" if model_name == "chatgpt-manual" else "ai"
            yield f"data: {json.dumps({'done': True, 'query_id': qrow.id, 'mode': mode, 'tokens': {'input': in_tok, 'output': out_tok, 'total': in_tok + out_tok}, 'duration_ms': dur, 'model': model_name})}\n\n"

        except Exception as e:
            error_msg = str(e)
            dur = int(time.time() * 1000) - t0
            try:
                log_e = models.Log(
                    user_display_name=user_display, section_id=section_id,
                    section_name=section_name, client_ip="stream",
                    has_file=bool(filename), filename=filename,
                    status="error", error_msg=error_msg, duration_ms=dur,
                )
                qrow = models.Query(
                    user_id=user_id, user_display_name=user_display,
                    section_id=section_id, section_name=section_name, section_icon=section_icon,
                    client_text=client_text or "", result="",
                    has_file=bool(filename), filename=filename,
                    status="error", error_msg=error_msg, duration_ms=dur,
                )
                db2.add(qrow); db2.add(log_e); db2.commit()
            except Exception:
                pass
            yield f"data: {json.dumps({'error': error_msg})}\n\n"
        finally:
            db2.close()

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ══════════════════════════════════════════════════════════════════════════════
# DOCX EXPORT
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/historial/{qid}/docx")
async def api_export_docx(req: Request, qid: int, db: Session = Depends(get_db)):
    u = _require(req)
    item = db.query(models.Query).filter(
        models.Query.id == qid, models.Query.user_id == u["id"]
    ).first()
    if not item:
        raise HTTPException(status_code=404)

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"{item.section_icon or ''} {item.section_name or 'Consulta'}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"{COMPANY_NAME}  ·  {item.user_display_name or u['display_name']}").bold = True
    if item.created_at:
        meta.add_run(f"  ·  {item.created_at.strftime('%d/%m/%Y %H:%M')}")

    doc.add_paragraph()

    if item.client_text:
        doc.add_heading("Consulta", level=2)
        doc.add_paragraph(item.client_text)
        doc.add_paragraph()

    if item.result:
        doc.add_heading("Resultado", level=2)
        clean = _strip_md(item.result)
        for para in clean.split("\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                if para.startswith("•"):
                    p.style = "List Bullet"
            else:
                doc.add_paragraph()

    doc.add_paragraph()
    footer_p = doc.add_paragraph(f"Generado por Gestor IA · {COMPANY_NAME}")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.color.rgb = RGBColor(0x88, 0x90, 0xA4)
    footer_p.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = re.sub(r"[^\w\-]", "_", item.section_name or "consulta")
    return StreamingResponse(
        iter([buf.read()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_name}_{qid}.docx"},
    )


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — API KEY STATUS
# ══════════════════════════════════════════════════════════════════════════════

_API_STATUS_TIMEOUT = 12  # seconds per API call

def _check_one_openai(key: str) -> dict:
    import httpx
    try:
        r = httpx.get(
            "https://api.openai.com/v1/models",
            headers={"Authorization": f"Bearer {key}"},
            timeout=_API_STATUS_TIMEOUT,
        )
        if r.status_code == 200:
            return {"status": "ok", "label": "Conectado"}
        return {"status": "error", "label": f"HTTP {r.status_code}"}
    except httpx.TimeoutException:
        return {"status": "error", "label": "Timeout"}
    except Exception as e:
        return {"status": "error", "label": str(e)[:80]}

def _check_one_gemini(key: str) -> dict:
    try:
        gclient = google_genai.Client(api_key=key)
        # Use a short timeout via httpx inside google client isn't easy,
        # so we wrap the list call with a thread timeout instead
        import threading, queue
        q: queue.Queue = queue.Queue()
        def _run():
            try:
                list(gclient.models.list())
                q.put(("ok", "Conectado"))
            except Exception as ex:
                q.put(("error", str(ex)[:80]))
        t = threading.Thread(target=_run, daemon=True)
        t.start()
        t.join(timeout=_API_STATUS_TIMEOUT)
        if t.is_alive():
            return {"status": "error", "label": "Timeout"}
        status, label = q.get()
        return {"status": status, "label": label}
    except Exception as e:
        return {"status": "error", "label": str(e)[:80]}

def _check_api_status_sync() -> dict:
    result = {}
    openai_key = os.getenv("OPENAI_API_KEY", "")
    gemini_key = os.getenv("GEMINI_API_KEY", "")

    futures = {}
    with ThreadPoolExecutor(max_workers=2) as pool:
        if not openai_key or openai_key.startswith("sk-..."):
            result["openai"] = {"status": "not_configured", "label": "Sin configurar"}
        else:
            futures["openai"] = pool.submit(_check_one_openai, openai_key)

        if not gemini_key:
            result["gemini"] = {"status": "not_configured", "label": "Sin configurar"}
        else:
            futures["gemini"] = pool.submit(_check_one_gemini, gemini_key)

        for key, fut in futures.items():
            result[key] = fut.result()

    return result


@app.get("/api/admin/api-status")
async def api_admin_api_status(req: Request):
    _require_admin(req)
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _check_api_status_sync)


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — EXTENDED METRICS
# ══════════════════════════════════════════════════════════════════════════════

MODEL_COST = {
    "gpt-4o":        (2.50, 10.00),
    "gpt-4o-mini":   (0.15,  0.60),
    "o3-mini":       (1.10,  4.40),
    "o1-mini":       (3.00, 12.00),
    "gpt-4-turbo":   (10.0, 30.00),
    "gpt-3.5-turbo": (0.50,  1.50),
}

def _cost_estimate(model: str, in_tok: int, out_tok: int) -> float:
    if model not in MODEL_COST:
        return 0.0
    in_rate, out_rate = MODEL_COST[model]
    return (in_tok / 1_000_000 * in_rate) + (out_tok / 1_000_000 * out_rate)

def _compute_metrics_sync() -> dict:
    import sqlite3 as _sqlite3
    conn = _sqlite3.connect("gestor_prompts.db", timeout=15)
    conn.row_factory = _sqlite3.Row
    cur = conn.cursor()
    cur.execute("PRAGMA journal_mode=WAL")

    def _q(sql):
        cur.execute(sql)
        rows = cur.fetchall()
        return [dict(r) for r in rows]

    queries_per_day = _q("""
        SELECT strftime('%Y-%m-%d', created_at) as day, COUNT(*) as total,
               SUM(CASE WHEN status='ok' THEN 1 ELSE 0 END) as ok_count
        FROM consultas
        WHERE created_at >= date('now', '-30 days')
        GROUP BY day ORDER BY day
    """)

    per_user_rows = _q("""
        SELECT user_id, user_display_name,
               COUNT(*) as total, SUM(CASE WHEN status='ok' THEN 1 ELSE 0 END) as ok_count,
               SUM(input_tokens + output_tokens) as total_tokens,
               AVG(duration_ms) as avg_ms,
               MAX(created_at) as last_query
        FROM consultas
        WHERE user_id IS NOT NULL
        GROUP BY user_id ORDER BY total DESC LIMIT 20
    """)
    per_user = [{**r, "avg_ms": round(r["avg_ms"] or 0)} for r in per_user_rows]

    per_section_rows = _q("""
        SELECT section_id, section_name, section_icon,
               COUNT(*) as total, SUM(CASE WHEN status='ok' THEN 1 ELSE 0 END) as ok_count,
               AVG(duration_ms) as avg_ms,
               SUM(input_tokens) as in_tok, SUM(output_tokens) as out_tok
        FROM consultas
        WHERE section_id IS NOT NULL
        GROUP BY section_id ORDER BY total DESC
    """)
    per_section = []
    for r in per_section_rows:
        per_section.append({**r,
            "avg_ms": round(r["avg_ms"] or 0),
            "error_rate": round((1 - (r["ok_count"] or 0) / max(r["total"], 1)) * 100, 1),
        })

    cur.execute("SELECT COUNT(DISTINCT user_id) FROM consultas WHERE date(created_at)=date('now')")
    active_today = cur.fetchone()[0] or 0

    cur.execute("SELECT COUNT(DISTINCT user_id) FROM consultas WHERE created_at >= date('now','-7 days')")
    active_week = cur.fetchone()[0] or 0

    cost_per_section = _q("""
        SELECT section_name as section, SUM(input_tokens) as in_tok, SUM(output_tokens) as out_tok
        FROM consultas WHERE status='ok' GROUP BY section_name
    """)

    tokens_per_day = _q("""
        SELECT strftime('%Y-%m-%d', created_at) as day,
               SUM(input_tokens) as in_tok, SUM(output_tokens) as out_tok
        FROM consultas WHERE created_at >= date('now','-7 days')
        GROUP BY day ORDER BY day
    """)

    conn.close()
    return {
        "queries_per_day": queries_per_day,
        "per_user": per_user,
        "per_section": per_section,
        "tokens_per_day": tokens_per_day,
        "cost_per_section": cost_per_section,
        "active_today": active_today,
        "active_week": active_week,
    }


@app.get("/api/admin/metrics")
async def api_admin_metrics(req: Request):
    _require_admin(req)
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _compute_metrics_sync)


@app.get("/api/historial")
async def api_historial(
    req: Request, page: int = 1, q: Optional[str] = None,
    section_id: Optional[int] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    status: Optional[str] = None,
    tag: Optional[str] = None,
    favorites: bool = False,
    db: Session = Depends(get_db),
):
    u = _require(req)
    from sqlalchemy import or_
    per_page = 20
    query = (db.query(models.Query).filter(models.Query.user_id == u["id"])
             .order_by(models.Query.created_at.desc()))
    if q and q.strip():
        term = f"%{q.strip()}%"
        query = query.filter(or_(
            models.Query.client_text.ilike(term),
            models.Query.section_name.ilike(term),
            models.Query.result.ilike(term),
            models.Query.tags.ilike(term),
        ))
    if section_id:
        query = query.filter(models.Query.section_id == section_id)
    if date_from:
        query = query.filter(models.Query.created_at >= datetime.strptime(date_from, "%Y-%m-%d"))
    if date_to:
        query = query.filter(models.Query.created_at < datetime.strptime(date_to, "%Y-%m-%d") + timedelta(days=1))
    if status:
        query = query.filter(models.Query.status == status)
    if tag:
        query = query.filter(models.Query.tags.ilike(f"%{tag}%"))
    if favorites:
        query = query.filter(models.Query.is_favorite == True)
    total = query.count()
    items = query.offset((page - 1) * per_page).limit(per_page).all()
    return {
        "items": [_query_dict(i) for i in items],
        "total": total, "pages": max(1, (total + per_page - 1) // per_page), "page": page,
        "sections": [{"id": s.id, "name": s.name, "icon": s.icon}
                     for s in _sections_for_user(db, u["id"])],
    }


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 1: Tags on historial queries
# ══════════════════════════════════════════════════════════════════════════════

@app.patch("/api/historial/{qid}/tags")
async def api_update_tags(req: Request, qid: int, body: Dict[str, Any], db: Session = Depends(get_db)):
    u = _require(req)
    item = db.query(models.Query).filter(models.Query.id == qid, models.Query.user_id == u["id"]).first()
    if not item:
        raise HTTPException(status_code=404)
    tags = body.get("tags", [])
    if not isinstance(tags, list):
        raise HTTPException(status_code=400, detail="tags debe ser una lista")
    item.tags = json.dumps([str(t).strip()[:50] for t in tags if str(t).strip()][:20])
    db.commit()
    return {"tags": json.loads(item.tags)}

@app.patch("/api/historial/{qid}/favorite")
async def api_toggle_favorite(req: Request, qid: int, db: Session = Depends(get_db)):
    u = _require(req)
    item = db.query(models.Query).filter(models.Query.id == qid, models.Query.user_id == u["id"]).first()
    if not item:
        raise HTTPException(status_code=404)
    item.is_favorite = not item.is_favorite
    db.commit()
    return {"is_favorite": item.is_favorite}


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 2: Retry a past query
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/historial/{qid}/retry")
async def api_retry(req: Request, qid: int, db: Session = Depends(get_db)):
    u = _require(req)
    orig = db.query(models.Query).filter(models.Query.id == qid, models.Query.user_id == u["id"]).first()
    if not orig:
        raise HTTPException(status_code=404)
    section = db.query(models.Section).filter(models.Section.id == orig.section_id).first()
    if not section:
        raise HTTPException(status_code=404, detail="Sección ya no existe")
    return {
        "section_id": orig.section_id,
        "section_name": orig.section_name,
        "client_text": orig.client_text or "",
        "had_file": orig.has_file,
        "query_id": orig.id,
        "filename": orig.filename,
    }


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 4: Admin — query detail
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/queries/{qid}")
async def api_admin_query_detail(req: Request, qid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    item = db.query(models.Query).filter(models.Query.id == qid).first()
    if not item:
        raise HTTPException(status_code=404)
    return _query_dict(item)


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 5: Usage limits helper
# ══════════════════════════════════════════════════════════════════════════════

def _check_limits(db: Session, user_id: int, section_id: int):
    """Raises HTTPException 429/403 if user has exceeded limits or lacks section access."""
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        return

    # Check section access: if any UserSection row exists for this section, it's restricted
    restricted = db.query(models.UserSection).filter(
        models.UserSection.section_id == section_id
    ).first()
    if restricted:
        allowed = db.query(models.UserSection).filter(
            models.UserSection.section_id == section_id,
            models.UserSection.user_id == user_id,
        ).first()
        if not allowed:
            raise HTTPException(status_code=403, detail="No tienes acceso a esta sección")

    today = datetime.now().date()
    month_start = today.replace(day=1)

    # Count today's queries for this user
    today_count = db.query(models.Query).filter(
        models.Query.user_id == user_id,
        models.Query.created_at >= datetime.combine(today, datetime.min.time()),
        models.Query.status == "ok",
    ).count()

    # Global daily limit
    if user.daily_limit and today_count >= user.daily_limit:
        raise HTTPException(status_code=429, detail=f"Límite diario alcanzado ({user.daily_limit} consultas/día)")

    # Monthly limit
    if user.monthly_limit:
        month_count = db.query(models.Query).filter(
            models.Query.user_id == user_id,
            models.Query.created_at >= datetime.combine(month_start, datetime.min.time()),
            models.Query.status == "ok",
        ).count()
        if month_count >= user.monthly_limit:
            raise HTTPException(status_code=429, detail=f"Límite mensual alcanzado ({user.monthly_limit} consultas/mes)")

    # Per-section daily limit
    us = db.query(models.UserSection).filter(
        models.UserSection.user_id == user_id,
        models.UserSection.section_id == section_id,
    ).first()
    if us and us.daily_limit:
        section_today = db.query(models.Query).filter(
            models.Query.user_id == user_id,
            models.Query.section_id == section_id,
            models.Query.created_at >= datetime.combine(today, datetime.min.time()),
            models.Query.status == "ok",
        ).count()
        if section_today >= us.daily_limit:
            raise HTTPException(status_code=429, detail=f"Límite diario para esta sección alcanzado ({us.daily_limit} consultas/día)")


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 6: Model comparison
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/compare/{sid}")
async def api_compare(
    req: Request, sid: int,
    client_text: Optional[str] = Form(None),
    model_b: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
):
    """Run same query against section's model AND model_b, return both results."""
    u = _require(req)
    section = db.query(models.Section).filter(models.Section.id == sid).first()
    if not section:
        raise HTTPException(status_code=404)
    if model_b not in AVAILABLE_MODELS:
        raise HTTPException(status_code=400, detail=f"Modelo desconocido: {model_b}")

    valid_files = [f for f in files if f and f.filename]
    prompt, messages, file_content, filename, all_names, is_image, image_mime, image_b64 = \
        await _prepare_request(section, client_text or "", valid_files, u)

    async def run_model(model_name: str, msgs) -> dict:
        t0 = int(time.time() * 1000)
        try:
            if model_name == "chatgpt-manual":
                return {"result": prompt, "model": model_name, "tokens": {}, "duration_ms": 0, "error": None}
            elif _is_gemini(model_name):
                gkey = os.getenv("GEMINI_API_KEY", "")
                if not gkey:
                    return {"result": None, "model": model_name, "tokens": {}, "duration_ms": 0, "error": "GEMINI_API_KEY no configurada"}
                gclient = google_genai.Client(api_key=gkey)
                cfg = genai_types.GenerateContentConfig(temperature=section.temperature, max_output_tokens=section.max_tokens)
                if is_image:
                    import PIL.Image
                    img_obj = PIL.Image.open(io.BytesIO(base64.b64decode(image_b64)))
                    resp = gclient.models.generate_content(model=model_name, contents=[prompt, img_obj], config=cfg)
                else:
                    resp = gclient.models.generate_content(model=model_name, contents=prompt, config=cfg)
                text = resp.text or ""
                dur = int(time.time() * 1000) - t0
                return {"result": text, "model": model_name, "tokens": {"input": 0, "output": 0, "total": 0}, "duration_ms": dur, "error": None}
            else:
                oai = _openai()
                resp = oai.chat.completions.create(
                    model=model_name, messages=msgs,
                    temperature=section.temperature, max_tokens=section.max_tokens,
                )
                text = resp.choices[0].message.content or ""
                usage = resp.usage
                dur = int(time.time() * 1000) - t0
                return {
                    "result": text, "model": model_name,
                    "tokens": {"input": usage.prompt_tokens, "output": usage.completion_tokens, "total": usage.total_tokens},
                    "duration_ms": dur, "error": None,
                }
        except Exception as e:
            return {"result": None, "model": model_name, "tokens": {}, "duration_ms": int(time.time() * 1000) - t0, "error": str(e)}

    import asyncio

    # Build messages for model_b using a copy of section with model_b set
    section_b = models.Section()
    section_b.__dict__.update(section.__dict__)
    section_b.model = model_b
    _, messages_b, _, _, _, _, _, _ = await _prepare_request(section_b, client_text or "", [], u)

    result_a, result_b = await asyncio.gather(
        run_model(section.model, messages),
        run_model(model_b, messages_b),
    )

    # Save comparison to historial
    try:
        q_rec = models.Query(
            user_id=u["id"],
            user_display_name=u.get("display_name", u["username"]),
            section_id=section.id,
            section_name=section.name,
            section_icon=section.icon,
            client_text=client_text or "",
            result=result_a.get("result") or "",
            has_file=bool(valid_files),
            filename=filename,
            input_tokens=result_a.get("tokens", {}).get("input", 0),
            output_tokens=result_a.get("tokens", {}).get("output", 0),
            duration_ms=result_a.get("duration_ms", 0),
            status="ok" if not result_a.get("error") else "error",
            error_msg=result_a.get("error"),
            is_comparison=True,
            model_b=model_b,
            result_b=result_b.get("result") or "",
        )
        db.add(q_rec)
        db.commit()
        db.refresh(q_rec)
        comparison_id = q_rec.id
    except Exception:
        comparison_id = None

    return {
        "model_a": result_a,
        "model_b": result_b,
        "section_name": section.name,
        "client_text": client_text or "",
        "query_id": comparison_id,
    }


# ══════════════════════════════════════════════════════════════════════════════
# USER LIMITS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/me/limits")
async def api_me_limits(req: Request, db: Session = Depends(get_db)):
    u = _require(req)
    user = db.query(models.User).filter(models.User.id == u["id"]).first()
    if not user:
        raise HTTPException(status_code=404)

    today = datetime.now().date()
    month_start = today.replace(day=1)

    today_count = db.query(models.Query).filter(
        models.Query.user_id == u["id"],
        models.Query.created_at >= datetime.combine(today, datetime.min.time()),
        models.Query.status == "ok",
        models.Query.is_comparison == False,
    ).count()

    month_count = db.query(models.Query).filter(
        models.Query.user_id == u["id"],
        models.Query.created_at >= datetime.combine(month_start, datetime.min.time()),
        models.Query.status == "ok",
        models.Query.is_comparison == False,
    ).count()

    return {
        "daily_used": today_count,
        "daily_limit": user.daily_limit,
        "monthly_used": month_count,
        "monthly_limit": user.monthly_limit,
    }


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN AUDIT LOG
# ══════════════════════════════════════════════════════════════════════════════

def _audit(db: Session, req: Request, action: str, resource_type: str = None,
           resource_id: int = None, resource_name: str = None, details: dict = None):
    try:
        db.add(models.AdminLog(
            action=action,
            resource_type=resource_type,
            resource_id=resource_id,
            resource_name=resource_name,
            details=json.dumps(details) if details else None,
            admin_ip=_ip(req),
        ))
        db.flush()
    except Exception as _audit_err:
        log.warning("Audit log failed: %s", _audit_err)

@app.get("/api/admin/audit")
async def api_admin_audit(req: Request, page: int = 1, db: Session = Depends(get_db)):
    _require_admin(req)
    per_page = 50
    q = db.query(models.AdminLog).order_by(models.AdminLog.created_at.desc())
    total = q.count()
    items = q.offset((page - 1) * per_page).limit(per_page).all()
    return {
        "items": [{
            "id": l.id, "action": l.action,
            "resource_type": l.resource_type, "resource_id": l.resource_id,
            "resource_name": l.resource_name, "details": l.details,
            "admin_ip": l.admin_ip,
            "created_at": l.created_at.isoformat() if l.created_at else None,
        } for l in items],
        "total": total,
        "pages": max(1, (total + per_page - 1) // per_page),
        "page": page,
    }


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN CLEANUP
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/cleanup/preview")
async def api_admin_cleanup_preview(req: Request, days: int = 30, db: Session = Depends(get_db)):
    _require_admin(req)
    cutoff = datetime.now() - timedelta(days=days)
    items = (db.query(models.Query)
             .filter(
                 models.Query.created_at < cutoff,
                 models.Query.folder_path != None,
                 models.Query.is_protected == False,
             )
             .order_by(models.Query.created_at.desc())
             .all())
    result = []
    total_size = 0
    for q in items:
        size = 0
        if q.folder_path and Path(q.folder_path).exists():
            for f in Path(q.folder_path).rglob("*"):
                if f.is_file():
                    size += f.stat().st_size
            total_size += size
        result.append({
            "id": q.id,
            "section_name": q.section_name,
            "user_display_name": q.user_display_name,
            "created_at": q.created_at.isoformat() if q.created_at else None,
            "has_file": q.has_file,
            "filename": q.filename,
            "folder_exists": q.folder_path is not None and Path(q.folder_path).exists(),
            "size_bytes": size,
            "is_protected": q.is_protected,
        })
    return {"items": result, "total": len(result), "total_size_bytes": total_size, "cutoff_days": days}


@app.post("/api/admin/cleanup/run")
async def api_admin_cleanup_run(req: Request, body: Dict[str, Any], db: Session = Depends(get_db)):
    _require_admin(req)
    days = int(body.get("days", 30))
    exclude_ids = set(body.get("exclude_ids", []))
    cutoff = datetime.now() - timedelta(days=days)

    items = (db.query(models.Query)
             .filter(
                 models.Query.created_at < cutoff,
                 models.Query.folder_path != None,
                 models.Query.is_protected == False,
             ).all())

    deleted_count = 0
    freed_bytes = 0
    for q in items:
        if q.id in exclude_ids:
            continue
        if q.folder_path and Path(q.folder_path).exists():
            import shutil
            size = sum(f.stat().st_size for f in Path(q.folder_path).rglob("*") if f.is_file())
            freed_bytes += size
            shutil.rmtree(q.folder_path, ignore_errors=True)
        q.folder_path = None
        deleted_count += 1

    _audit(db, req, "cleanup.run", "cleanup", None, None,
           {"days": days, "deleted": deleted_count, "freed_bytes": freed_bytes})
    db.commit()
    log.info("Cleanup: deleted files for %d queries, freed %.1f MB", deleted_count, freed_bytes / 1_000_000)
    return {"deleted": deleted_count, "freed_bytes": freed_bytes}


@app.patch("/api/admin/queries/{qid}/protect")
async def api_admin_toggle_protect(req: Request, qid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    q = db.query(models.Query).filter(models.Query.id == qid).first()
    if not q:
        raise HTTPException(status_code=404)
    q.is_protected = not q.is_protected
    db.commit()
    return {"id": q.id, "is_protected": q.is_protected}




# ══════════════════════════════════════════════════════════════════════════════
# SECTION VERSIONS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/sections/{sid}/versions")
async def api_section_versions(req: Request, sid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    versions = (db.query(models.SectionVersion)
                .filter(models.SectionVersion.section_id == sid)
                .order_by(models.SectionVersion.created_at.desc())
                .limit(50)
                .all())
    return [{
        "id": v.id, "section_name": v.section_name,
        "prompt": v.prompt, "model": v.model,
        "temperature": v.temperature, "max_tokens": v.max_tokens,
        "changed_by_ip": v.changed_by_ip,
        "created_at": v.created_at.isoformat() if v.created_at else None,
    } for v in versions]


@app.post("/api/admin/sections/{sid}/versions/{vid}/restore")
async def api_section_restore(req: Request, sid: int, vid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    s = db.query(models.Section).filter(models.Section.id == sid).first()
    if not s:
        raise HTTPException(status_code=404, detail="Sección no encontrada")
    v = db.query(models.SectionVersion).filter(
        models.SectionVersion.id == vid,
        models.SectionVersion.section_id == sid,
    ).first()
    if not v:
        raise HTTPException(status_code=404, detail="Versión no encontrada")
    # Save current state as a new version before restoring
    db.add(models.SectionVersion(
        section_id=s.id, section_name=s.name,
        prompt=s.prompt, model=s.model,
        temperature=s.temperature, max_tokens=s.max_tokens,
        changed_by_ip=_ip(req),
    ))
    s.prompt = v.prompt
    if v.model:
        s.model = v.model
    if v.temperature is not None:
        s.temperature = v.temperature
    if v.max_tokens is not None:
        s.max_tokens = v.max_tokens
    _audit(db, req, "section.restore", "section", s.id, s.name,
           {"restored_version_id": vid})
    db.commit()
    return _sec_dict(s)


# ══════════════════════════════════════════════════════════════════════════════
# PASSWORD RESET
# ══════════════════════════════════════════════════════════════════════════════

import secrets as _secrets

@app.post("/api/admin/users/{uid}/reset-token")
async def api_generate_reset_token(req: Request, uid: int, db: Session = Depends(get_db)):
    _require_admin(req)
    user = db.query(models.User).filter(models.User.id == uid).first()
    if not user:
        raise HTTPException(status_code=404)
    # Invalidate previous unused tokens for this user
    db.query(models.PasswordResetToken).filter(
        models.PasswordResetToken.user_id == uid,
        models.PasswordResetToken.used == False,
    ).delete()
    token = _secrets.token_urlsafe(32)
    expires = datetime.now() + timedelta(hours=24)
    db.add(models.PasswordResetToken(token=token, user_id=uid, expires_at=expires))
    db.commit()
    log.info("Reset token generated for user '%s' by admin from %s", user.username, _ip(req))
    return {"token": token, "expires_at": expires.isoformat(), "username": user.username}


@app.get("/api/reset-password/validate")
async def api_validate_reset_token(token: str, db: Session = Depends(get_db)):
    rec = db.query(models.PasswordResetToken).filter(
        models.PasswordResetToken.token == token,
        models.PasswordResetToken.used == False,
    ).first()
    if not rec or rec.expires_at < datetime.now():
        raise HTTPException(status_code=400, detail="Token inválido o caducado")
    user = db.query(models.User).filter(models.User.id == rec.user_id).first()
    if not user:
        raise HTTPException(status_code=400, detail="Usuario no encontrado")
    return {"valid": True, "username": user.username, "display_name": user.display_name}


@app.post("/api/reset-password")
async def api_reset_password(body: Dict[str, Any], db: Session = Depends(get_db)):
    token = body.get("token", "").strip()
    new_password = body.get("password", "")
    if not token or not new_password or len(new_password) < 6:
        raise HTTPException(status_code=400, detail="Token y contraseña (mínimo 6 caracteres) requeridos")
    rec = db.query(models.PasswordResetToken).filter(
        models.PasswordResetToken.token == token,
        models.PasswordResetToken.used == False,
    ).first()
    if not rec or rec.expires_at < datetime.now():
        raise HTTPException(status_code=400, detail="Token inválido o caducado")
    user = db.query(models.User).filter(models.User.id == rec.user_id).first()
    if not user:
        raise HTTPException(status_code=400, detail="Usuario no encontrado")
    user.password_hash = await _hash(new_password)
    rec.used = True
    db.commit()
    log.info("Password reset for user '%s'", user.username)
    return {"ok": True, "username": user.username}


# ══════════════════════════════════════════════════════════════════════════════
# QUERY CACHE (in-memory, text-only queries)
# ══════════════════════════════════════════════════════════════════════════════

import hashlib as _hashlib

_query_cache: dict = {}
_CACHE_TTL = 3600  # 1 hour


def _cache_key(user_id: int, section_id: int, text: str) -> str:
    h = _hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest()
    return f"{user_id}:{section_id}:{h}"


def _cache_get(key: str) -> Optional[dict]:
    entry = _query_cache.get(key)
    if entry and (time.time() - entry["_ts"]) < _CACHE_TTL:
        return entry
    if key in _query_cache:
        del _query_cache[key]
    return None


def _cache_set(key: str, value: dict):
    _query_cache[key] = {**value, "_ts": time.time()}
    now = time.time()
    expired = [k for k, v in _query_cache.items() if (now - v["_ts"]) >= _CACHE_TTL]
    for k in expired:
        del _query_cache[k]
    if len(_query_cache) > 500:
        oldest = min(_query_cache, key=lambda k: _query_cache[k]["_ts"])
        del _query_cache[oldest]


# ══════════════════════════════════════════════════════════════════════════════
# STATIC + SPA
# ══════════════════════════════════════════════════════════════════════════════

app.mount("/data", StaticFiles(directory="data"), name="data")
app.mount("/asserts", StaticFiles(directory="asserts"), name="asserts")

_spa_dist = Path("frontend/dist")
if _spa_dist.exists():
    app.mount("/assets", StaticFiles(directory=str(_spa_dist / "assets")), name="spa-assets")

@app.get("/{full_path:path}")
async def serve_spa(full_path: str):
    index = _spa_dist / "index.html"
    if index.exists():
        return FileResponse(str(index))
    return JSONResponse(
        {"detail": "Frontend no compilado. Ejecuta: cd frontend && npm run build"},
        status_code=503,
    )
